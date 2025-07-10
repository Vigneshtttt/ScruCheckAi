from flask import Flask, request, render_template, send_file
import os
import pytesseract
from PIL import Image
import requests
import io
import markdown2 
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_PARAGRAPH_ALIGNMENT
from docx.shared import Pt, Inches  
import matplotlib
import fitz  
import re
import matplotlib.pyplot as plt

app = Flask(__name__)
UPLOAD_FOLDER = 'uploads'
os.makedirs(UPLOAD_FOLDER, exist_ok=True)

MISTRAL_API_KEY = "EcUefyGHLhCczHMJCvspkOdU6lJSe73O"

# Header/Footer filter configuration
HEADER_FOOTER_KEYWORDS = [
    "EnggTree.com", "Downloaded from", "Question Paper Code", 
    "Reg.No", "Page", "30236", "B.E./B.Tech"
]
HEADER_REGION = 0.10  # Top 10% of page
FOOTER_REGION = 0.85  # Bottom 15% of page

# === Text Cleaning ===
def clean_extracted_text(text):
    """Remove repetitive headers/footers and redundant newlines"""
    # Remove header/footer lines
    lines = text.split('\n')
    cleaned_lines = []
    for line in lines:
        if not any(keyword.lower() in line.lower() for keyword in HEADER_FOOTER_KEYWORDS):
            cleaned_lines.append(line)
    
    # Join and clean extra spaces
    cleaned_text = '\n'.join(cleaned_lines)
    cleaned_text = re.sub(r'\n{3,}', '\n\n', cleaned_text)  # Reduce multiple newlines
    cleaned_text = re.sub(r' {2,}', ' ', cleaned_text)      # Reduce multiple spaces
    return cleaned_text.strip()

# === Text Extraction ===
def extract_text(file_path):
    ext = os.path.splitext(file_path)[-1].lower()

    if ext == '.txt':
        with open(file_path, 'r', encoding='utf-8') as f:
            return clean_extracted_text(f.read())

    elif ext == '.docx':
        doc = Document(file_path)
        text = "\n".join([para.text for para in doc.paragraphs])
        return clean_extracted_text(text)

    elif ext == '.pdf':
        text = extract_text_from_pdf_or_ocr(file_path)
        return clean_extracted_text(text)

    else:
        return "Unsupported file type."

def extract_text_from_pdf_or_ocr(pdf_path):
    try:
        doc = fitz.open(pdf_path)
    except Exception as e:
        return f"Error opening PDF: {str(e)}"

    extracted_text = ""
    
    for page in doc:
        try:
            # Get text blocks with position information
            blocks = page.get_text("blocks")
            page_text = ""
            
            for block in blocks:
                # Unpack block: (x0, y0, x1, y1, text, block_no, block_type)
                y0 = block[1]
                y1 = block[3]
                text = block[4]
                
                # Skip header/footer regions
                if y0 < page.rect.height * HEADER_REGION or y1 > page.rect.height * FOOTER_REGION:
                    continue
                
                page_text += text + "\n"
            
            # Fallback to OCR if no valid text found in content area
            if not page_text.strip():
                pix = page.get_pixmap()
                img = Image.frombytes("RGB", [pix.width, pix.height], pix.samples)
                ocr_text = pytesseract.image_to_string(img)
                page_text = ocr_text
                
            extracted_text += page_text
            
        except Exception as e:
            extracted_text += f"\n[Error processing page: {str(e)}]"
    
    return extracted_text

# === Mistral API Call ===
def analyze_sample(sample_text, prev_text, syllabus_text):
    url = "https://api.mistral.ai/v1/chat/completions"
    headers = {
        "Authorization": f"Bearer {MISTRAL_API_KEY}",
        "Content-Type": "application/json"
    }
    prompt = f"""
You are an academic scrutiny AI designed to evaluate university exam question papers according to standardized criteria. Given the SAMPLE QUESTION PAPER, PREVIOUS QUESTION PAPER, and SYLLABUS, perform an in-depth analysis and provide a structured Markdown report based on the 10 scrutiny criteria below.

### 🎯 Scrutiny Criteria

1. **Format compliance**: Does the paper follow the required structure? (Part A: 10×2=20, Part B: 5×16=80 marks, proper numbering, section labels)
2. **Regulation & course compliance**: Is the paper aligned with the correct regulation, semester, and course code as per the syllabus?(the question paper should have all these details)
3. **Syllabus alignment**: Are all questions within syllabus scope? Identify any out-of-syllabus content.
4. **Bloom’s Taxonomy classification**: Classify each question's cognitive level (Remember, Understand, Apply, Analyze, Evaluate, Create).
5. **Mark distribution & time balance**: Are marks fairly distributed across difficulty and time? Flag overly lengthy or light questions.
6. **Grammar and clarity check**: Are all questions grammatically correct and unambiguous?
7. **Use of diagrams/symbols**: Are diagrams or equations used appropriately, and are they necessary?
8. **Permitted aids check**: Are references to calculators, tables, or handbooks in line with what is allowed?
9. **Repetition check**: Are any questions repeated from the previous paper? Highlight them.()
10. **Figure naming clarity**: Are figures correctly labeled (e.g., Fig.1(a)) and referenced properly?

---

### 📄 Report Format (in Markdown)

#### 1. Format Compliance  
- ✅ Result: YES  ❌ Result: NO
- Reason: ...

#### 2. Regulation & Course Check  
- ✅ Result: YES  ❌ Result: NO
- Reason: ...

#### 3. Syllabus Alignment  
- ✅ Result: YES  ❌ Result: NO
- Reason: ...  
- Out-of-syllabus Questions:  
  - Q13: "Explain XYZ" → Not found in syllabus

#### ✅ 4. Bloom’s Taxonomy & COs  
- Result: ✅ YES / ❌ NO  
- Reason: ...   
| Question No | Bloom Level      |
|-------------|------------------|
| Q1          | Remember         |
| Q2          | Apply            |
| Q3          | ...              |
| Q4          | ...              |
| Q5          | ...              |
| Q6          | ...              |
| Q7          | ...              |
| Q8          | ...              |
| Q9          | ...              |
| Q10         | ...              |
| Q11 (a)     | ...              |
| Q11 (b)     | ...              |
| Q12 (a)     | ...              |
| Q12 (b)     | ...              |
| ...         | ...              |


#### 5. Mark Distribution & Time  
- ✅ Result: YES  ❌ Result: NO
- Reason: ...

#### 6. Grammar & Clarity  
- ✅ Result: YES  ❌ Result: NO
- Notable Issues:  
  - Q6: Ambiguous verb usage  
  - Q10: Missing punctuation

#### 7. Diagrams/Symbols  
- ✅ Result: YES  ❌ Result: NO
- Comments: ...

#### 8. Permitted Aids  
- ✅ Result: YES  ❌ Result: NO
- Reason: ...

#### 9. Repetition Check  
- Inter-paper Result: ...
- Inter-paper Repeated Questions: ...
- Intra-paper Result: ...
- Intra-paper Repeated Questions: ...



#### 10. Figure Naming Clarity  
- ✅ Result: YES  ❌ Result: NO
- Issues:  
  - "Refer to diagram below" → No label

---

Use only the available content. Skip incomplete sections with a warning. Ensure answers are based strictly on evidence from the papers and syllabus.

---

### SAMPLE QUESTION PAPER:
{sample_text[:10000]}

### PREVIOUS YEAR QUESTION PAPER:
{prev_text[:10000]}

### SYLLABUS:
{syllabus_text[:5000]}

"""

    payload = {
        "model": "mistral-medium",
        "messages": [
            {"role": "system", "content": "You analyze academic exam papers."},
            {"role": "user", "content": prompt}
        ],
        "temperature": 0.3
    }
    try:
        response = requests.post(url, headers=headers, json=payload, timeout=60)
        return response.json()["choices"][0]["message"]["content"]
    except Exception as e:
        return f"Error calling Mistral API: {str(e)}"

# === Report Generation ===
def generate_chart(data_dict, title):
    labels = list(data_dict.keys())
    sizes = [int(val.strip('%')) for val in data_dict.values()]
    fig, ax = plt.subplots()
    ax.pie(sizes, labels=labels, autopct='%1.1f%%', startangle=90)
    ax.axis('equal')
    plt.title(title)
    img_stream = io.BytesIO()
    plt.savefig(img_stream, format='png', bbox_inches='tight')
    plt.close(fig)
    img_stream.seek(0)
    return img_stream

def save_report_to_docx(report_text):
    doc = Document()
    doc.add_heading("📘 Exam Question Paper Analysis", level=1)

    summary_charts = {}
    sections = report_text.split("### ")
    for section in sections:
        if not section.strip():
            continue
        lines = section.strip().splitlines()
        heading = lines[0].strip()
        doc.add_heading(heading, level=2)

        if "|" in "\n".join(lines):
            rows = [line for line in lines[1:] if "|" in line and "---" not in line]
            if rows:
                headers = [h.strip() for h in rows[0].split("|")[1:-1]]
                table = doc.add_table(rows=1, cols=len(headers))
                table.style = "Table Grid"
                hdr_cells = table.rows[0].cells
                for i, h in enumerate(headers):
                    hdr_cells[i].text = h
                for row in rows[1:]:
                    cells = [c.strip() for c in row.split("|")[1:-1]]
                    row_cells = table.add_row().cells
                    for i, val in enumerate(cells):
                        row_cells[i].text = val
        else:
            for line in lines[1:]:
                clean = line.strip()
                if heading.startswith("4. Statistical Summary") and ":" in clean and "%" in clean:
                    key, val = clean.split(":")
                    section_name = lines[0].split("Distribution")[0].strip("• ").strip()
                    summary_charts.setdefault(section_name, {})[key.strip()] = val.strip()
                    doc.add_paragraph(clean)
                elif clean.startswith(("-", "*", "•")):
                    doc.add_paragraph(clean[1:].strip(), style="List Bullet")
                elif clean and clean[0].isdigit() and clean[1:3] == ". ":
                    doc.add_paragraph(clean[3:].strip(), style="List Number")
                elif clean:
                    doc.add_paragraph(clean)

    if summary_charts:
        doc.add_page_break()
        doc.add_heading("📊 Visual Summary Charts", level=1)
        for chart_title, data in summary_charts.items():
            doc.add_heading(chart_title, level=2)
            img = generate_chart(data, chart_title)
            doc.add_picture(img, width=Inches(5))
            doc.paragraphs[-1].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    report_io = io.BytesIO()
    doc.save(report_io)
    report_io.seek(0)
    return report_io

# === Flask Routes ===
@app.route('/', methods=['GET', 'POST'])
def index():
    analysis_result = None

    if request.method == 'POST':
        sample_file = request.files.get('sample_file')
        prev_file = request.files.get('prev_file')  # Optional
        syllabus_file = request.files.get('syllabus_file')

        # Mandatory file checks
        if not sample_file or not syllabus_file:
            analysis_result = markdown2.markdown(
                "**Error:** Please upload both Sample Question Paper and Syllabus. Previous Year Paper is optional."
            )
            return render_template('index.html', result=analysis_result)

        # Save uploaded files
        sample_path = os.path.join(UPLOAD_FOLDER, sample_file.filename)
        sample_file.save(sample_path)
        sample_text = extract_text(sample_path)

        syllabus_path = os.path.join(UPLOAD_FOLDER, syllabus_file.filename)
        syllabus_file.save(syllabus_path)
        syllabus_text = extract_text(syllabus_path)

        # If previous paper is uploaded, process it
        if prev_file and prev_file.filename:
            prev_path = os.path.join(UPLOAD_FOLDER, prev_file.filename)
            prev_file.save(prev_path)
            prev_text = extract_text(prev_path)
        else:
            prev_text = "Not Provided"

        # Content validation
        if len(sample_text) < 100 or len(syllabus_text) < 50:
            analysis_result = markdown2.markdown(
                "**Error:** Insufficient content in required files. "
                "Please ensure Sample Paper and Syllabus contain valid content."
            )
        else:
            raw_result = analyze_sample(sample_text, prev_text, syllabus_text)
            analysis_result = markdown2.markdown(raw_result, extras=["tables"])  

            # Save raw markdown for download
            with open("uploads/temp_analysis.txt", "w", encoding="utf-8") as f:
                f.write(raw_result)

    return render_template('index.html', result=analysis_result)



@app.route('/download', methods=['POST'])
def download_report():
    try:
        with open("uploads/temp_analysis.txt", "r", encoding="utf-8") as f:
            report_text = f.read()
        report_file = save_report_to_docx(report_text)
        return send_file(report_file, download_name="exam_analysis_report.docx", as_attachment=True)
    except Exception as e:
        return f"Error preparing download: {str(e)}"

if __name__ == '__main__':
    app.run(debug=True)